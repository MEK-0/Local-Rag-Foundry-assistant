import docx
from typing import List

from src.parser.base import BaseParser, KnowledgeNode, NodeType

# python-docx exposes heading level via paragraph.style.name, e.g. "Heading 1", "Heading 2"
HEADING_STYLE_PREFIX = "Heading"


class DocxDocumentParser(BaseParser):
    def parse_to_tree(self, file_path: str, doc_id: str) -> List[KnowledgeNode]:
        """
        Converts a .docx file into a KnowledgeNode tree.
        - Paragraphs styled "Heading N" become SECTION nodes (level = N),
          building the same parent-chain logic as the PDF parser.
        - Regular paragraphs attach to the current heading as PARAGRAPH nodes.
        - Tables (doc.tables) are extracted as markdown TABLE nodes, attached
          to whichever heading precedes them in document order.
        """
        doc = docx.Document(file_path)
        nodes: List[KnowledgeNode] = []
        heading_stack: List[KnowledgeNode] = []
        order_counter = 0

        # docx separates paragraphs and tables into different collections;
        # walk doc.element.body in document order so tables land under the
        # correct heading instead of all being appended at the end
        table_iter = iter(doc.tables)
        table_by_element = {t._tbl: t for t in doc.tables}

        for block in doc.element.body:
            tag = block.tag.split("}")[-1]

            if tag == "p":
                para = next((p for p in doc.paragraphs if p._p is block), None)
                if para is None or not para.text.strip():
                    continue

                style_name = para.style.name if para.style else ""
                if style_name.startswith(HEADING_STYLE_PREFIX):
                    level = self._heading_level(style_name)
                    order_counter += 1

                    while heading_stack and heading_stack[-1].level >= level:
                        heading_stack.pop()
                    parent = heading_stack[-1] if heading_stack else None

                    node = KnowledgeNode(
                        doc_id=doc_id,
                        parent_id=parent.id if parent else None,
                        type=NodeType.SECTION,
                        level=level,
                        content=para.text.strip(),
                        order=order_counter,
                    )
                    node.heading_path = self._path_for(parent, node.content)
                    heading_stack.append(node)
                    nodes.append(node)
                else:
                    order_counter += 1
                    parent = heading_stack[-1] if heading_stack else None
                    nodes.append(
                        KnowledgeNode(
                            doc_id=doc_id,
                            parent_id=parent.id if parent else None,
                            type=NodeType.PARAGRAPH,
                            level=(parent.level + 1) if parent else 1,
                            heading_path=self._path_for(parent),
                            content=para.text.strip(),
                            order=order_counter,
                        )
                    )

            elif tag == "tbl":
                table_obj = table_by_element.get(block)
                if table_obj is None:
                    continue
                md_table = self._table_to_markdown(table_obj)
                if not md_table:
                    continue
                order_counter += 1
                parent = heading_stack[-1] if heading_stack else None
                nodes.append(
                    KnowledgeNode(
                        doc_id=doc_id,
                        parent_id=parent.id if parent else None,
                        type=NodeType.TABLE,
                        level=(parent.level + 1) if parent else 1,
                        heading_path=self._path_for(parent),
                        content=md_table,
                        order=order_counter,
                        metadata={"rows": len(table_obj.rows)},
                    )
                )

        return nodes

    @staticmethod
    def _heading_level(style_name: str) -> int:
        # "Heading 1" -> 1, "Heading 2" -> 2, fallback to 1 if unparseable
        digits = "".join(c for c in style_name if c.isdigit())
        return int(digits) if digits else 1

    @staticmethod
    def _path_for(parent, current_title: str = None) -> str:
        base = parent.heading_path if parent else ""
        if parent and not base:
            base = parent.content
        if current_title:
            return f"{base} > {current_title}" if base else current_title
        return base

    @staticmethod
    def _table_to_markdown(table_obj) -> str:
        rows = [[cell.text.strip() for cell in row.cells] for row in table_obj.rows]
        rows = [r for r in rows if any(c for c in r)]
        if not rows:
            return ""
        header, *body = rows
        md = "| " + " | ".join(header) + " |\n"
        md += "| " + " | ".join(["---"] * len(header)) + " |\n"
        for row in body:
            md += "| " + " | ".join(c.replace("\n", " ") for c in row) + " |\n"
        return md